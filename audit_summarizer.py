import streamlit as st
import openai
from docx import Document
from PyPDF2 import PdfReader
from io import BytesIO

# -------- CONFIG --------
st.set_page_config(page_title="Audit Board Minutes Summarizer", layout="centered")
st.title("📄 Audit Board Minutes Summarizer")

# -------- API KEY --------
openai.api_key = st.text_input("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", type="password")

# -------- File Upload --------
uploaded_file = st.file_uploader("📎 Upload a .docx or .pdf file with board minutes", type=["docx", "pdf"])

# -------- Enhanced Prompt Builder --------
def build_prompt(text):
    return f"""
You are a senior internal audit assistant.

Read the board meeting minutes below and generate a **high-quality summary table** for audit documentation.

For each key point discussed in the meeting, extract:

1. **Date of Meeting**
2. **Attendees** (if mentioned)
3. **Key Points** – Use concise but detailed bullet points. Mention any:
   - Decisions made
   - Issues raised
   - Actions assigned (to whom and when)
   - Financial, legal, or operational implications
4. **Audit Impact** – Clearly assess the relevance to audit, including:
   - "Control deficiency identified"
   - "Risk of non-compliance"
   - "Opportunity for process improvement"
   - "No audit implication"

⚠️ Use formal and professional tone. This summary will be used in an **internal audit file**.

Return your response as a **well-formatted Markdown table only**, without any extra explanation.

Board Minutes:
{text}
"""

# -------- Text Extractors --------
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

# -------- AI Summary --------
def summarize_minutes(text):
    prompt = build_prompt(text)
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

# -------- Export to DOCX --------
def generate_docx(markdown_table):
    doc = Document()
    doc.add_heading("Board Minutes Summary for Audit", level=1)

    # Markdown → Table
    lines = [line for line in markdown_table.split("\n") if "|" in line]
    rows = [line.split("|")[1:-1] for line in lines if "---" not in line]

    if not rows:
        return None

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"

    for i, header in enumerate(rows[0]):
        table.rows[0].cells[i].text = header.strip()

    for row_data in rows[1:]:
        row = table.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = cell.strip()

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# -------- Main Logic --------
if uploaded_file and openai.api_key:
    with st.spinner("🔎 Summarizing..."):
        try:
            if uploaded_file.name.endswith(".docx"):
                minutes_text = extract_text_from_docx(uploaded_file)
            elif uploaded_file.name.endswith(".pdf"):
                minutes_text = extract_text_from_pdf(uploaded_file)
            else:
                st.error("Unsupported file type.")
                st.stop()

            summary = summarize_minutes(minutes_text)
            st.success("✅ Summary complete!")

            st.markdown("### 🧾 Summary Output:")
            st.markdown(summary)

            docx_file = generate_docx(summary)
            if docx_file:
                st.download_button("⬇️ Download Word Summary", docx_file, file_name="audit_summary.docx")
            else:
                st.warning("Could not generate downloadable summary.")
        except Exception as e:
            st.error(f"Something went wrong: {e}")
